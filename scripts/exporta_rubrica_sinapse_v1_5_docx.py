from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pymongo

from pymongo import MongoClient

with MongoClient("mongodb://admin:admin123@localhost:27017/?authSource=admin") as cliente:
    db = cliente["rubricas"]
    colecao = db["rubrica_sinapse"]
    rubrica = colecao.find_one({"versao": "v1.5", "status": "ativa"})

# Criar documento Word
doc = Document()

# Título
titulo = doc.add_heading("Rubrica SINAPSE IA – Versão v1.5", level=1)
titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph(
    "Instrumento avaliativo com base em SAEB, PISA OCDE e princípios neuropsicopedagógicos e metodológicos atualizados.",
    style="Normal"
)

# Percorrer dimensões
for dim in rubrica["dimensoes"][0]["dimensoes"]:
    doc.add_heading(f'\n🧠 {dim["dimensao"]}', level=2)
    doc.add_paragraph(f'Fonte: {dim.get("fonte", "Não especificada")}', style="Intense Quote")
    if dim.get("nota"):
        doc.add_paragraph(f'Nota: {dim["nota"]}', style="Intense Quote")

    # Tabela
    tabela = doc.add_table(rows=1, cols=3)
    tabela.style = "Light List Accent 1"
    hdr_cells = tabela.rows[0].cells
    hdr_cells[0].text = "Nível"
    hdr_cells[1].text = "Descritor"
    hdr_cells[2].text = "Exemplo Prático"

    exemplos = dim.get("exemplos", [])

    for nivel in dim["niveis"]:
        linha = tabela.add_row().cells
        linha[0].text = f"{nivel['nota']} – {nivel['nome']}"
        linha[1].text = nivel["descricao"]

        # Exibir 3 exemplos em linhas separadas dentro da mesma célula
        if exemplos:
            texto_exemplos = "\n".join(exemplos[:3])
            paragrafo = linha[2].paragraphs[0]
            for exemplo in texto_exemplos.split("\n"):
                run = paragrafo.add_run(f"{exemplo}\n")
                run.font.size = Pt(10)

# Página de referências
doc.add_page_break()
doc.add_heading("Referências (Estilo ABNT)", level=2)

for ref in rubrica["dimensoes"][0].get("referencias", []):
    autor = ref["autor"]
    ano = ref["ano"]
    titulo = f"**{ref['titulo']}**"
    local = ref.get("local", "")
    editora = ref.get("editora", "")
    revista = ref.get("revista", "")
    volume = ref.get("volume", "")
    paginas = ref.get("paginas", "")

    if revista:
        texto_ref = f"{autor}. {ref['titulo']}. {revista}, v. {volume}, p. {paginas}, {ano}."
    else:
        texto_ref = f"{autor}. {ref['titulo']}. {local}: {editora}, {ano}."

    doc.add_paragraph(texto_ref)

# Salvar arquivo
caminho = "/home/neirivon/Downloads/rubrica_sinapse_ia_v1.5.docx"
doc.save(caminho)
print(f"📄 Arquivo '{caminho}' gerado com sucesso.")

