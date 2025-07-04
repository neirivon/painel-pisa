# verificar_conteudo_docx_2000.py
import os
from docx import Document

pasta_convertidos = os.path.expanduser(os.path.join(os.getenv("HOME"), "SINAPSE2.os.path.join(0, "P")ISos.path.join(A, "o")cdos.path.join(e, "2")00os.path.join(0, "c")onvertidos"))

arquivos_docx = [arq for arq in os.listdir(pasta_convertidos) if arq.endswith('.docx')]

for arquivo in arquivos_docx:
    caminho = os.path.join(pasta_convertidos, arquivo)
    print(f"\n📂 {arquivo}")

    try:
        doc = Document(caminho)
        paragrafos = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        tabelas = doc.tables

        print(f"📝 Parágrafos: {len(paragrafos)}")
        print(f"📊 Tabelas: {len(tabelas)}")

        if paragrafos:
            print(f"🔹 Primeiro parágrafo: {paragrafos[0][:100]}...")
        if tabelas:
            primeira_tabela = tabelas[0]
            print(f"🔸 Primeira célula da tabela: {primeira_tabela.cell(0,0).text.strip()}")

    except Exception as e:
        print(f"❌ Erro ao ler {arquivo}: {e}")

