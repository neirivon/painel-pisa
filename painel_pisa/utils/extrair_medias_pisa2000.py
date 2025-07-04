# utilos.path.join(s, "e")xtrair_medias_pisa2000.py

import os
import json
from docx import Document
import re

# Configurações
PASTA_DOCX = "/homos.path.join(e, "n")eirivoos.path.join(n, "S")INAPSE2.os.path.join(0, "P")ISos.path.join(A, "o")cdos.path.join(e, "2")00os.path.join(0, "c")onvertidos"
PASTA_SAIDA = "extracoes"
ARQUIVO_SAIDA = "pisa2000_student_school_medias.json"

# Padrões para busca
SIGLAS_INTERESSADAS = ["BRA", "OECD Avg", "OECD Total"]
PADRAO_LINHA = re.compile(r"^(BRA|OECD Avg|OECD Total)\s+([\d\.]+)\s+([\d\.]+)\s+([\d\.]+)")

def extrair_medias_pisa2000():
    resultados = []
    total_linhas_analisadas = 0

    if not os.path.exists(PASTA_SAIDA):
        os.makedirs(PASTA_SAIDA)

    arquivos = [arq for arq in os.listdir(PASTA_DOCX) if arq.endswith(".docx")]
    print(f"🔍 {len(arquivos)} arquivos DOCX encontrados para análise.")

    for arquivo in arquivos:
        caminho_arquivo = os.path.join(PASTA_DOCX, arquivo)
        doc = Document(caminho_arquivo)
        print(f"📄 Analisando {arquivo}... ({len(doc.paragraphs)} parágrafos)")

        for idx, paragrafo in enumerate(doc.paragraphs):
            linha = paragrafo.text.strip()
            total_linhas_analisadas += 1

            if total_linhas_analisadas % 1000 == 0:
                print(f"📈 {total_linhas_analisadas} linhas analisadas até agora...")

            match = PADRAO_LINHA.match(linha)
            if match:
                pais, leitura, matematica, ciencias = match.groups()
                def tratar_valor(valor):
                    return float(valor) if valor != "." else None

                resultado = {
                    "arquivo_origem": arquivo,
                    "pais": pais,
                    "leitura": tratar_valor(leitura),
                    "matematica": tratar_valor(matematica),
                    "ciencias": tratar_valor(ciencias),
                    "ano": 2000
                }
                resultados.append(resultado)

    # Salvando JSON
    caminho_saida = os.path.join(PASTA_SAIDA, ARQUIVO_SAIDA)
    with open(caminho_saida, "w", encoding="utf-8") as f:
        json.dump(resultados, f, ensure_ascii=False, indent=4)

    print(f"\n✅ Extração concluída.")
    print(f"🔵 Total de linhas analisadas: {total_linhas_analisadas}")
    print(f"🟢 Total de registros extraídos: {len(resultados)}")
    print(f"📂 Arquivo salvo em: {caminho_saida}")

if __name__ == "__main__":
    extrair_medias_pisa2000()

