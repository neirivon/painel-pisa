# utilos.path.join(s, "e")xtrair_medias_pisa2000_avancado.py

import os
import json
import re
from docx import Document
from collections import defaultdict

# Diretórios
PASTA_DOCX = "/homos.path.join(e, "n")eirivoos.path.join(n, "S")INAPSE2.os.path.join(0, "P")ISos.path.join(A, "o")cdos.path.join(e, "2")00os.path.join(0, "c")onvertidos"
PASTA_SAIDA = "extracoes"
ARQUIVO_SAIDA = "pisa2000_student_medias_avancado.json"

# Siglas que nos interessam
SIGLAS_INTERESSADAS = ["BRA", "OECD Avg", "OECD Total"]

# Expressão regular para capturar linhas com dados
PADRAO_LINHA = re.compile(r"^(BRA|OECD Avg|OECD Total)\s+([\d,.]+)\s+([\d,.]+)\s+([\d,.]+)")

# Função de limpeza de números
def limpar_numero(valor):
    valor = valor.replace(",", ".").replace(" ", "").strip()
    try:
        return float(valor)
    except ValueError:
        return None

# Função principal de extração
def extrair_medias_pisa2000_avancado():
    resultados = []

    if not os.path.exists(PASTA_SAIDA):
        os.makedirs(PASTA_SAIDA)

    arquivos = [arq for arq in os.listdir(PASTA_DOCX) if arq.endswith(".docx")]
    print(f"🔍 {len(arquivos)} arquivos DOCX encontrados para análise.")

    for arquivo in arquivos:
        caminho_arquivo = os.path.join(PASTA_DOCX, arquivo)
        doc = Document(caminho_arquivo)
        print(f"📄 Analisando {arquivo}... ({len(doc.paragraphs)} parágrafos)")

        # Primeiro: ler parágrafos comuns
        for idx, paragrafo in enumerate(doc.paragraphs):
            if idx % 1000 == 0 and idx > 0:
                print(f"📈 {idx} parágrafos analisados...")
            linha = paragrafo.text.strip()
            match = PADRAO_LINHA.match(linha)
            if match:
                pais, leitura, matematica, ciencias = match.groups()
                resultados.append({
                    "arquivo_origem": arquivo,
                    "pais": pais,
                    "leitura": limpar_numero(leitura),
                    "matematica": limpar_numero(matematica),
                    "ciencias": limpar_numero(ciencias),
                    "ano": 2000
                })

        # Segundo: ler tabelas
        print(f"📚 Analisando {len(doc.tables)} tabelas no documento...")
        for t_idx, tabela in enumerate(doc.tables):
            for r_idx, linha in enumerate(tabela.rows):
                textos = [celula.text.strip() for celula in linha.cells]
                if not textos or len(textos) < 4:
                    continue
                linha_completa = " ".join(textos)
                match = PADRAO_LINHA.match(linha_completa)
                if match:
                    pais, leitura, matematica, ciencias = match.groups()
                    resultados.append({
                        "arquivo_origem": arquivo,
                        "pais": pais,
                        "leitura": limpar_numero(leitura),
                        "matematica": limpar_numero(matematica),
                        "ciencias": limpar_numero(ciencias),
                        "ano": 2000
                    })
            if (t_idx + 1) % 10 == 0:
                print(f"📊 {t_idx+1} tabelas processadas...")

    # Salvar o JSON final
    caminho_saida = os.path.join(PASTA_SAIDA, ARQUIVO_SAIDA)
    with open(caminho_saida, "w", encoding="utf-8") as f:
        json.dump(resultados, f, ensure_ascii=False, indent=4)

    print(f"✅ Extração científica avançada concluída. {len(resultados)} registros salvos em {caminho_saida}")

if __name__ == "__main__":
    extrair_medias_pisa2000_avancado()

