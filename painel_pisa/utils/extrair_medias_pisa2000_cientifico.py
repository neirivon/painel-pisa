# utilos.path.join(s, "e")xtrair_medias_pisa2000_cientifico.py

import os
import json
import re
from docx import Document

# Diretórios e configurações iniciais
PASTA_DOCX = "/homos.path.join(e, "n")eirivoos.path.join(n, "S")INAPSE2.os.path.join(0, "P")ISos.path.join(A, "o")cdos.path.join(e, "2")00os.path.join(0, "c")onvertidos"
PASTA_SAIDA = "extracoes"
ARQUIVO_SAIDA = "pisa2000_student_medias_validado.json"

SIGLAS_INTERESSADAS = ["BRA", "OECD Avg", "OECD Total"]
LIMITE_MIN = 100
LIMITE_MAX = 700

# Expressão regular super-rigorosa para detectar linha de média
PADRAO_MEDIA = re.compile(r"^(BRA|OECD Avg|OECD Total)\s+(\d+(\.\d+)?)\s+(\d+(\.\d+)?)\s+(\d+(\.\d+)?)$")

# Função principal

def extrair_medias_pisa2000_cientifico():
    resultados = []

    if not os.path.exists(PASTA_SAIDA):
        os.makedirs(PASTA_SAIDA)

    arquivos = [arq for arq in os.listdir(PASTA_DOCX) if arq.startswith("Student") and arq.endswith(".docx")]

    print(f"🔍 {len(arquivos)} arquivos Student_compendium encontrados para análise.")

    for arquivo in arquivos:
        caminho = os.path.join(PASTA_DOCX, arquivo)
        doc = Document(caminho)

        print(f"\n📄 Analisando {arquivo}... ({len(doc.paragraphs)} parágrafos)")

        linhas_analisadas = 0

        for paragrafo in doc.paragraphs:
            linha = paragrafo.text.strip()
            linhas_analisadas += 1

            if linhas_analisadas % 1000 == 0:
                print(f"📈 {linhas_analisadas} linhas analisadas...")

            match = PADRAO_MEDIA.match(linha)
            if match:
                pais, leitura, _, matematica, _, ciencias, _ = match.groups()
                leitura = float(leitura)
                matematica = float(matematica)
                ciencias = float(ciencias)

                if all(LIMITE_MIN <= v <= LIMITE_MAX for v in [leitura, matematica, ciencias]):
                    resultado = {
                        "arquivo_origem": arquivo,
                        "pais": pais,
                        "leitura": leitura,
                        "matematica": matematica,
                        "ciencias": ciencias,
                        "ano": 2000
                    }
                    resultados.append(resultado)
                else:
                    print(f"⚠️ Média descartada fora dos limites: {linha}")

    # Salvar resultado
    caminho_saida = os.path.join(PASTA_SAIDA, ARQUIVO_SAIDA)
    with open(caminho_saida, "w", encoding="utf-8") as f:
        json.dump(resultados, f, ensure_ascii=False, indent=4)

    print(f"\n✅ Extração científica concluída. {len(resultados)} registros válidos salvos em {caminho_saida}")

if __name__ == "__main__":
    extrair_medias_pisa2000_cientifico()

